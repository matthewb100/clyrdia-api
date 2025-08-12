"""
Document processing service for various file formats
"""
import hashlib
import io
from typing import Optional, Tuple
from fastapi import UploadFile, HTTPException
import PyPDF2
from docx import Document
import aiofiles
import structlog

logger = structlog.get_logger(__name__)


class DocumentService:
    """Service for processing various document formats"""
    
    def __init__(self):
        self.supported_formats = {
            "application/pdf": self._extract_pdf_text,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": self._extract_docx_text,
            "text/plain": self._extract_text_content,
        }
    
    async def process_document(
        self, 
        file: UploadFile,
        max_size: int = 10485760  # 10MB default
    ) -> Tuple[str, str]:
        """
        Process uploaded document and extract text content
        
        Returns:
            Tuple of (extracted_text, document_hash)
        """
        try:
            # Validate file size
            if file.size and file.size > max_size:
                raise HTTPException(
                    status_code=413,
                    detail=f"File size {file.size} bytes exceeds maximum allowed size {max_size} bytes"
                )
            
            # Validate file type
            if file.content_type not in self.supported_formats:
                raise HTTPException(
                    status_code=400,
                    detail=f"Unsupported file type: {file.content_type}. Supported types: {list(self.supported_formats.keys())}"
                )
            
            # Read file content
            content = await file.read()
            
            # Extract text based on file type
            extractor = self.supported_formats[file.content_type]
            extracted_text = await extractor(content)
            
            # Generate document hash
            document_hash = self._generate_hash(content)
            
            logger.info(
                "Document processed successfully",
                filename=file.filename,
                content_type=file.content_type,
                size=len(content),
                text_length=len(extracted_text),
                hash=document_hash
            )
            
            return extracted_text, document_hash
            
        except HTTPException:
            raise
        except Exception as e:
            logger.error("Document processing failed", filename=file.filename, error=str(e))
            raise HTTPException(
                status_code=500,
                detail=f"Failed to process document: {str(e)}"
            )
    
    async def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF content"""
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_content = ""
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content += page_text + "\n"
            
            if not text_content.strip():
                raise ValueError("No text content found in PDF")
            
            return text_content.strip()
            
        except Exception as e:
            logger.error("PDF text extraction failed", error=str(e))
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    async def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX content"""
        try:
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)
            
            text_content = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content += cell.text + "\n"
            
            if not text_content.strip():
                raise ValueError("No text content found in DOCX")
            
            return text_content.strip()
            
        except Exception as e:
            logger.error("DOCX text extraction failed", error=str(e))
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
    
    async def _extract_text_content(self, content: bytes) -> str:
        """Extract text from plain text content"""
        try:
            text_content = content.decode('utf-8')
            
            if not text_content.strip():
                raise ValueError("No text content found")
            
            return text_content.strip()
            
        except UnicodeDecodeError:
            # Try other encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    text_content = content.decode(encoding)
                    if text_content.strip():
                        return text_content.strip()
                except UnicodeDecodeError:
                    continue
            
            raise ValueError("Failed to decode text content with any supported encoding")
    
    def _generate_hash(self, content: bytes) -> str:
        """Generate SHA-256 hash of document content"""
        return hashlib.sha256(content).hexdigest()
    
    async def validate_document_content(self, text: str) -> bool:
        """Validate extracted document content"""
        if not text or not text.strip():
            return False
        
        # Check minimum content length
        if len(text.strip()) < 50:
            return False
        
        # Check for common document indicators
        document_indicators = [
            "contract", "agreement", "terms", "conditions", "clause",
            "party", "parties", "effective date", "termination",
            "liability", "indemnification", "governing law"
        ]
        
        text_lower = text.lower()
        indicator_count = sum(1 for indicator in document_indicators if indicator in text_lower)
        
        # Require at least 2 document indicators
        return indicator_count >= 2
    
    async def sanitize_text(self, text: str) -> str:
        """Sanitize extracted text content"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # Clean individual lines
            cleaned_line = ' '.join(line.split())
            if cleaned_line:
                cleaned_lines.append(cleaned_line)
        
        # Join lines with reasonable spacing
        cleaned_text = '\n'.join(cleaned_lines)
        
        # Remove excessive newlines
        while '\n\n\n' in cleaned_text:
            cleaned_text = cleaned_text.replace('\n\n\n', '\n\n')
        
        return cleaned_text.strip()
    
    async def extract_metadata(self, content: bytes, filename: str, content_type: str) -> dict:
        """Extract basic metadata from document"""
        metadata = {
            "filename": filename,
            "content_type": content_type,
            "size_bytes": len(content),
            "hash": self._generate_hash(content)
        }
        
        try:
            if content_type == "application/pdf":
                pdf_file = io.BytesIO(content)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                
                metadata.update({
                    "page_count": len(pdf_reader.pages),
                    "pdf_version": pdf_reader.pdf_header,
                    "is_encrypted": pdf_reader.is_encrypted
                })
                
                # Try to extract document info
                if pdf_reader.metadata:
                    metadata["pdf_metadata"] = {
                        "title": pdf_reader.metadata.get('/Title'),
                        "author": pdf_reader.metadata.get('/Author'),
                        "subject": pdf_reader.metadata.get('/Subject'),
                        "creator": pdf_reader.metadata.get('/Creator'),
                        "producer": pdf_reader.metadata.get('/Producer'),
                        "creation_date": pdf_reader.metadata.get('/CreationDate'),
                        "modification_date": pdf_reader.metadata.get('/ModDate')
                    }
            
            elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                docx_file = io.BytesIO(content)
                doc = Document(docx_file)
                
                metadata.update({
                    "paragraph_count": len(doc.paragraphs),
                    "table_count": len(doc.tables),
                    "section_count": len(doc.sections)
                })
                
                # Extract document properties
                if doc.core_properties:
                    metadata["docx_properties"] = {
                        "title": doc.core_properties.title,
                        "author": doc.core_properties.author,
                        "subject": doc.core_properties.subject,
                        "created": doc.core_properties.created,
                        "modified": doc.core_properties.modified,
                        "last_modified_by": doc.core_properties.last_modified_by,
                        "revision": doc.core_properties.revision
                    }
            
        except Exception as e:
            logger.warning("Failed to extract metadata", error=str(e))
            metadata["metadata_extraction_error"] = str(e)
        
        return metadata


# Global service instance
document_service = DocumentService() 